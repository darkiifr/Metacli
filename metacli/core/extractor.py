"""Metadata extraction functionality for various file types."""

import os
import mimetypes
import hashlib
import functools
import logging
from pathlib import Path
from typing import Dict, Any, Optional, List, Union
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
import time
from collections import OrderedDict
from enum import Enum
import gc
import psutil

# Configure logger for metadata extraction
logger = logging.getLogger(__name__)


class MemoryMonitor:
    """Monitor system memory usage and trigger cleanup when needed."""
    
    def __init__(self, warning_threshold: float = 0.8, critical_threshold: float = 0.9):
        """
        Initialize memory monitor.
        
        Args:
            warning_threshold: Memory usage percentage to trigger warnings (0.0-1.0)
            critical_threshold: Memory usage percentage to trigger aggressive cleanup (0.0-1.0)
        """
        self.warning_threshold = warning_threshold
        self.critical_threshold = critical_threshold
        self._last_check = 0
        self._check_interval = 5.0  # Check every 5 seconds
        self._lock = threading.RLock()
        
    def get_memory_usage(self) -> Dict[str, float]:
        """Get current memory usage statistics."""
        try:
            memory = psutil.virtual_memory()
            return {
                'percent': memory.percent / 100.0,
                'available_gb': memory.available / (1024**3),
                'used_gb': memory.used / (1024**3),
                'total_gb': memory.total / (1024**3)
            }
        except Exception as e:
            logger.warning(f"Failed to get memory usage: {e}")
            return {'percent': 0.0, 'available_gb': 0.0, 'used_gb': 0.0, 'total_gb': 0.0}
    
    def should_trigger_gc(self) -> bool:
        """Check if garbage collection should be triggered."""
        with self._lock:
            current_time = time.time()
            if current_time - self._last_check < self._check_interval:
                return False
            
            self._last_check = current_time
            memory_info = self.get_memory_usage()
            memory_percent = memory_info['percent']
            
            if memory_percent >= self.critical_threshold:
                logger.warning(f"Critical memory usage: {memory_percent:.1%} - triggering aggressive cleanup")
                return True
            elif memory_percent >= self.warning_threshold:
                logger.info(f"High memory usage: {memory_percent:.1%} - triggering garbage collection")
                return True
            
            return False
    
    def trigger_cleanup(self, aggressive: bool = False) -> Dict[str, Any]:
        """Trigger memory cleanup and return statistics."""
        before_memory = self.get_memory_usage()
        
        # Standard garbage collection
        collected = gc.collect()
        
        if aggressive:
            # More aggressive cleanup
            for generation in range(3):
                gc.collect(generation)
            
            # Clear metadata cache if available
            try:
                MetadataExtractor._cache.clear()
                logger.info("Cleared metadata cache during aggressive cleanup")
            except Exception as e:
                logger.debug(f"Could not clear cache: {e}")
        
        after_memory = self.get_memory_usage()
        
        cleanup_stats = {
            'objects_collected': collected,
            'memory_before_percent': before_memory['percent'],
            'memory_after_percent': after_memory['percent'],
            'memory_freed_mb': (before_memory['used_gb'] - after_memory['used_gb']) * 1024,
            'aggressive': aggressive
        }
        
        logger.info(f"Memory cleanup completed: {cleanup_stats}")
        return cleanup_stats


class ExtractionError(Exception):
    """Base exception for metadata extraction errors."""
    pass


class FileAccessError(ExtractionError):
    """Exception raised when file cannot be accessed."""
    pass


class UnsupportedFileTypeError(ExtractionError):
    """Exception raised when file type is not supported."""
    pass


class CorruptedFileError(ExtractionError):
    """Exception raised when file appears to be corrupted."""
    pass


class DependencyMissingError(ExtractionError):
    """Exception raised when required dependency is missing."""
    pass


class ErrorSeverity(Enum):
    """Severity levels for extraction errors."""
    LOW = "low"          # Non-critical errors, partial extraction possible
    MEDIUM = "medium"    # Significant errors, basic metadata only
    HIGH = "high"        # Critical errors, extraction failed
    CRITICAL = "critical" # System-level errors, may affect other operations

try:
    from PIL import Image
    from PIL.ExifTags import TAGS
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

try:
    import mutagen
    from mutagen.id3 import ID3NoHeaderError
    MUTAGEN_AVAILABLE = True
except ImportError:
    MUTAGEN_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class MetadataCache:
    """Thread-safe LRU cache for metadata results with memory management."""
    
    def __init__(self, max_size: int = 500, ttl: int = 3600):  # Reduced default size
        self.max_size = max_size
        self.ttl = ttl  # Time to live in seconds
        self._cache = OrderedDict()
        self._timestamps = {}
        self._lock = threading.RLock()
        self._memory_threshold = 100 * 1024 * 1024  # 100MB memory limit
        self._current_memory = 0
    
    def _estimate_size(self, value: Dict[str, Any]) -> int:
        """Estimate memory size of cached value."""
        try:
            import sys
            return sys.getsizeof(value)
        except:
            return 1024  # Default estimate
    
    def get(self, key: str) -> Optional[Dict[str, Any]]:
        with self._lock:
            if key in self._cache:
                # Check if entry is still valid
                if time.time() - self._timestamps[key] < self.ttl:
                    # Move to end (most recently used)
                    self._cache.move_to_end(key)
                    return self._cache[key].copy()  # Return a copy to avoid reference issues
                else:
                    # Entry expired, remove it
                    old_value = self._cache[key]
                    self._current_memory -= self._estimate_size(old_value)
                    del self._cache[key]
                    del self._timestamps[key]
            return None
    
    def set(self, key: str, value: Dict[str, Any]) -> None:
        with self._lock:
            value_size = self._estimate_size(value)
            
            # If updating existing key, remove old memory usage
            if key in self._cache:
                old_value = self._cache[key]
                self._current_memory -= self._estimate_size(old_value)
                del self._cache[key]
                del self._timestamps[key]
            
            # Remove oldest entries if cache is full or memory limit exceeded
            while (len(self._cache) >= self.max_size or 
                   self._current_memory + value_size > self._memory_threshold):
                if not self._cache:
                    break
                oldest_key = next(iter(self._cache))
                old_value = self._cache[oldest_key]
                self._current_memory -= self._estimate_size(old_value)
                del self._cache[oldest_key]
                del self._timestamps[oldest_key]
            
            # Store a copy to avoid reference issues
            self._cache[key] = value.copy()
            self._timestamps[key] = time.time()
            self._current_memory += value_size
    
    def clear(self) -> None:
        with self._lock:
            self._cache.clear()
            self._timestamps.clear()
            self._current_memory = 0
    
    def size(self) -> int:
        with self._lock:
            return len(self._cache)
    
    def memory_usage(self) -> int:
        """Get current memory usage estimate."""
        with self._lock:
            return self._current_memory


class MetadataExtractor:
    """Extract metadata from various file types with performance optimizations."""
    
    # Class-level cache for metadata results
    _cache = MetadataCache(max_size=500, ttl=3600)
    _cache_lock = threading.RLock()
    
    # Class-level memory monitor
    _memory_monitor = MemoryMonitor(warning_threshold=0.75, critical_threshold=0.85)
    _memory_monitor_lock = threading.RLock()
    
    # Pre-compiled supported types for faster lookup
    _SUPPORTED_TYPES = {
        'image': frozenset(['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif']),
        'audio': frozenset(['.mp3', '.flac', '.ogg', '.m4a', '.wav', '.wma']),
        'video': frozenset(['.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm']),
        'document': frozenset(['.pdf', '.docx', '.doc', '.txt', '.rtf']),
        'archive': frozenset(['.zip', '.rar', '.7z', '.tar', '.gz'])
    }
    
    # Flattened extension set for O(1) lookup
    _ALL_EXTENSIONS = frozenset().union(*_SUPPORTED_TYPES.values())
    
    # Class-level thread pool for reuse
    _thread_pool = None
    _thread_pool_lock = threading.RLock()
    
    def __init__(self, enable_cache: bool = True, max_workers: int = 4):
        """Initialize the metadata extractor.
        
        Args:
            enable_cache: Whether to enable metadata caching
            max_workers: Maximum number of threads for parallel processing
        """
        self.enable_cache = enable_cache
        self.max_workers = min(max_workers, 4)  # Limit max workers to prevent resource exhaustion
        self.supported_types = {k: list(v) for k, v in self._SUPPORTED_TYPES.items()}
        self._ensure_thread_pool()
    
    @classmethod
    def get_memory_stats(cls) -> Dict[str, Any]:
        """Get current memory usage statistics."""
        with cls._memory_monitor_lock:
            memory_info = cls._memory_monitor.get_memory_usage()
            cache_info = {
                'cache_size': cls._cache.size(),
                'cache_memory_mb': cls._cache.memory_usage() / (1024 * 1024)
            }
            return {**memory_info, **cache_info}
    
    @classmethod
    def check_and_cleanup_memory(cls, force: bool = False) -> Optional[Dict[str, Any]]:
        """Check memory usage and trigger cleanup if needed."""
        with cls._memory_monitor_lock:
            if force or cls._memory_monitor.should_trigger_gc():
                memory_info = cls._memory_monitor.get_memory_usage()
                aggressive = memory_info['percent'] >= cls._memory_monitor.critical_threshold
                return cls._memory_monitor.trigger_cleanup(aggressive=aggressive)
            return None
    
    @classmethod
    def force_memory_cleanup(cls) -> Dict[str, Any]:
        """Force aggressive memory cleanup."""
        with cls._memory_monitor_lock:
            return cls._memory_monitor.trigger_cleanup(aggressive=True)
    
    def get_file_type(self, file_path: Path) -> str:
        """Determine the file type category with optimized lookup."""
        suffix = file_path.suffix.lower()
        for file_type, extensions in self._SUPPORTED_TYPES.items():
            if suffix in extensions:
                return file_type
        return 'unknown'
    
    def _get_cache_key(self, file_path: Path) -> str:
        """Generate a cache key based on file path and modification time."""
        try:
            stat = file_path.stat()
            content = f"{file_path.absolute()}:{stat.st_mtime}:{stat.st_size}"
            return hashlib.md5(content.encode()).hexdigest()
        except (OSError, IOError):
            return str(file_path.absolute())
    
    def _get_cached_metadata(self, cache_key: str) -> Optional[Dict[str, Any]]:
        """Retrieve cached metadata if available."""
        if not self.enable_cache:
            return None
        
        return self._cache.get(cache_key)
    
    def _cache_metadata(self, cache_key: str, metadata: Dict[str, Any]) -> None:
        """Cache metadata result."""
        if not self.enable_cache:
            return
        
        self._cache.set(cache_key, metadata)
    
    def extract_basic_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract basic file system metadata with enhanced error handling."""
        try:
            stat = file_path.stat()
            
            metadata = {
                'filename': file_path.name,
                'filepath': str(file_path.absolute()),
                'size': stat.st_size,
                'size_human': self._format_size(stat.st_size),
                'size_kb': round(stat.st_size / 1024, 2),
                'size_mb': round(stat.st_size / (1024 * 1024), 2),
                'created': datetime.fromtimestamp(stat.st_ctime).isoformat(),
                'modified': datetime.fromtimestamp(stat.st_mtime).isoformat(),
                'accessed': datetime.fromtimestamp(stat.st_atime).isoformat(),
                'extension': file_path.suffix.lower(),
                'mime_type': mimetypes.guess_type(str(file_path))[0],
                'file_type': self.get_file_type(file_path),
                'is_hidden': file_path.name.startswith('.'),
                'parent_directory': str(file_path.parent),
                'stem': file_path.stem
            }
            
            # Add file permissions (cross-platform)
            try:
                metadata['permissions'] = oct(stat.st_mode)[-3:]
                metadata['is_readable'] = os.access(file_path, os.R_OK)
                metadata['is_writable'] = os.access(file_path, os.W_OK)
                metadata['is_executable'] = os.access(file_path, os.X_OK)
            except (OSError, AttributeError):
                pass
            
            return metadata
            
        except (OSError, IOError) as e:
            return {
                'filename': file_path.name if file_path else 'unknown',
                'filepath': str(file_path) if file_path else 'unknown',
                'error': f'Failed to extract basic metadata: {str(e)}'
            }
    
    def extract_image_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from image files with enhanced information and lazy loading."""
        metadata = {}
        
        if not PIL_AVAILABLE:
            metadata['error'] = 'PIL/Pillow not available for image metadata extraction'
            return metadata
        
        try:
            # Check file size first - skip detailed extraction for very large files
            file_size = file_path.stat().st_size
            if file_size > 50 * 1024 * 1024:  # 50MB limit
                metadata['size_warning'] = 'Large file - limited metadata extraction'
                # Only extract basic info for large files
                with Image.open(file_path) as img:
                    metadata.update({
                        'width': img.width,
                        'height': img.height,
                        'mode': img.mode,
                        'format': img.format
                    })
                return metadata
            
            # Use context manager for proper resource cleanup
            with Image.open(file_path) as img:
                # Basic image properties with enhanced information
                metadata.update({
                    'width': img.width,
                    'height': img.height,
                    'mode': img.mode,
                    'format': img.format,
                    'has_transparency': img.mode in ('RGBA', 'LA') or 'transparency' in img.info,
                    'pixel_count': img.width * img.height,
                    'aspect_ratio': round(img.width / img.height, 2) if img.height > 0 else 0,
                    'megapixels': round((img.width * img.height) / 1000000, 2),
                    'color_depth': len(img.getbands()) if hasattr(img, 'getbands') else None
                })
                
                # Add image info
                if hasattr(img, 'info') and img.info:
                    info_data = {}
                    for key, value in img.info.items():
                        if isinstance(value, (str, int, float, bool)):
                            info_data[key] = value
                        elif isinstance(value, (bytes, bytearray)):
                            info_data[key] = f'<binary_data_{len(value)}_bytes>'
                    if info_data:
                        metadata['image_info'] = info_data
                
                # Extract EXIF data with reduced processing for performance
                exif_dict = img.getexif()
                if exif_dict:
                    exif_data = {}
                    # Reduced essential tags for better performance
                    essential_tag_ids = [306, 271, 272, 274]  # DateTime, Make, Model, Orientation
                    for tag_id, value in exif_dict.items():
                        if tag_id in essential_tag_ids:
                            try:
                                tag = TAGS.get(tag_id, f'Tag_{tag_id}')
                                # Convert complex types to strings safely
                                if isinstance(value, (bytes, bytearray)):
                                    exif_data[tag] = f'<binary_data_{len(value)}_bytes>'
                                else:
                                    exif_data[tag] = str(value)[:100]  # Reduced limit
                            except Exception:
                                continue  # Skip problematic tags
                    
                    if exif_data:
                        metadata['exif'] = exif_data
                
        except (IOError, OSError) as e:
            metadata['error'] = f'Failed to open image file: {str(e)}'
        except Exception as e:
            metadata['error'] = f'Failed to extract image metadata: {str(e)}'
        
        return metadata
    
    def extract_audio_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from audio files with enhanced information and performance optimization."""
        metadata = {}
        
        if not MUTAGEN_AVAILABLE:
            metadata['error'] = 'Mutagen not available for audio metadata extraction'
            return metadata
        
        try:
            # Check file size first - skip detailed extraction for very large files
            file_size = file_path.stat().st_size
            if file_size > 100 * 1024 * 1024:  # 100MB limit for audio
                metadata['size_warning'] = 'Large audio file - limited metadata extraction'
                # Only extract basic info for large files
                audio_file = mutagen.File(file_path)
                if audio_file and hasattr(audio_file, 'info'):
                    info = audio_file.info
                    metadata.update({
                        'duration': getattr(info, 'length', 0),
                        'bitrate': getattr(info, 'bitrate', 0),
                        'sample_rate': getattr(info, 'sample_rate', 0),
                        'channels': getattr(info, 'channels', 0)
                    })
                return metadata
            
            audio_file = mutagen.File(file_path)
            if audio_file is None:
                metadata['error'] = 'Unsupported audio format'
                return metadata
            
            # Basic audio info with enhanced details
            if hasattr(audio_file, 'info'):
                info = audio_file.info
                duration = getattr(info, 'length', 0)
                bitrate = getattr(info, 'bitrate', 0)
                sample_rate = getattr(info, 'sample_rate', 0)
                channels = getattr(info, 'channels', 0)
                
                metadata.update({
                    'duration': duration,
                    'duration_human': self._format_duration(duration),
                    'duration_minutes': round(duration / 60, 2) if duration > 0 else 0,
                    'bitrate': bitrate,
                    'bitrate_kbps': f'{bitrate} kbps' if bitrate > 0 else 'Unknown',
                    'sample_rate': sample_rate,
                    'sample_rate_khz': f'{sample_rate / 1000:.1f} kHz' if sample_rate > 0 else 'Unknown',
                    'channels': channels,
                    'channel_mode': 'Stereo' if channels == 2 else 'Mono' if channels == 1 else f'{channels} channels',
                    'quality': 'High' if bitrate >= 320 else 'Medium' if bitrate >= 128 else 'Low' if bitrate > 0 else 'Unknown'
                })
                
                # Calculate estimated file size for different qualities
                if duration > 0:
                    metadata['estimated_sizes'] = {
                        '128_kbps': f'{round((128 * duration) / 8 / 1024, 2)} MB',
                        '192_kbps': f'{round((192 * duration) / 8 / 1024, 2)} MB',
                        '320_kbps': f'{round((320 * duration) / 8 / 1024, 2)} MB'
                    }
            
            # Enhanced tag extraction with reduced processing for performance
            if audio_file.tags:
                tags = {}
                # Limit tag processing for performance
                tag_count = 0
                max_tags = 20  # Limit number of tags processed
                
                for key, value in audio_file.tags.items():
                    if tag_count >= max_tags:
                        break
                    try:
                        if isinstance(value, list):
                            tags[key] = [str(v)[:50] for v in value if v]  # Reduced length
                        else:
                            tags[key] = str(value)[:50]  # Reduced length
                        tag_count += 1
                    except (TypeError, AttributeError):
                        continue
                
                if tags:
                    metadata['tags'] = tags
                
                # Reduced common tag mappings for better performance
                common_tags = {
                    'title': ['TIT2', 'TITLE', '\xa9nam'],
                    'artist': ['TPE1', 'ARTIST', '\xa9ART'],
                    'album': ['TALB', 'ALBUM', '\xa9alb'],
                    'date': ['TDRC', 'DATE', '\xa9day'],
                    'genre': ['TCON', 'GENRE', '\xa9gen']
                }
                
                for common_name, possible_keys in common_tags.items():
                    for key in possible_keys:
                        if key in tags and tags[key]:
                            value = tags[key][0] if isinstance(tags[key], list) else tags[key]
                            if value and str(value).strip():
                                metadata[common_name] = str(value).strip()[:50]  # Reduced length
                                break
        
        except Exception as e:
            metadata['error'] = f'Failed to extract audio metadata: {str(e)}'
        
        return metadata
    
    def extract_video_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from video files with enhanced information and performance optimization."""
        metadata = {}
        
        # For now, use mutagen for basic video metadata
        if not MUTAGEN_AVAILABLE:
            metadata['error'] = 'Mutagen not available for video metadata extraction'
            return metadata
        
        try:
            # Check file size first - skip detailed extraction for very large files
            file_size = file_path.stat().st_size
            if file_size > 500 * 1024 * 1024:  # 500MB limit for video
                metadata['size_warning'] = 'Large video file - limited metadata extraction'
                # Only extract basic info for large files
                video_file = mutagen.File(file_path)
                if video_file and hasattr(video_file, 'info'):
                    info = video_file.info
                    metadata.update({
                        'duration': getattr(info, 'length', 0),
                        'bitrate': getattr(info, 'bitrate', 0),
                        'file_size_mb': round(file_size / (1024 * 1024), 2)
                    })
                return metadata
            
            video_file = mutagen.File(file_path)
            if video_file is None:
                metadata['error'] = 'Unsupported video format'
                return metadata
            
            if hasattr(video_file, 'info'):
                info = video_file.info
                duration = getattr(info, 'length', 0)
                bitrate = getattr(info, 'bitrate', 0)
                
                metadata.update({
                    'duration': duration,
                    'duration_human': self._format_duration(duration),
                    'duration_minutes': round(duration / 60, 2) if duration > 0 else 0,
                    'duration_hours': round(duration / 3600, 2) if duration > 0 else 0,
                    'bitrate': bitrate,
                    'bitrate_kbps': f'{bitrate} kbps' if bitrate > 0 else 'Unknown',
                    'file_size_mb': round(file_size / (1024 * 1024), 2),
                    'file_size_gb': round(file_size / (1024 * 1024 * 1024), 3)
                })
                
                # Add video-specific properties if available
                if hasattr(info, 'width') and hasattr(info, 'height'):
                    width = getattr(info, 'width', 0)
                    height = getattr(info, 'height', 0)
                    metadata.update({
                        'width': width,
                        'height': height,
                        'resolution': f'{width}x{height}' if width and height else 'Unknown',
                        'aspect_ratio': round(width / height, 2) if height > 0 else 0,
                        'total_pixels': width * height if width and height else 0
                    })
                    
                    # Determine video quality category
                    if width >= 3840:  # 4K
                        metadata['quality_category'] = '4K Ultra HD'
                    elif width >= 1920:  # 1080p
                        metadata['quality_category'] = 'Full HD (1080p)'
                    elif width >= 1280:  # 720p
                        metadata['quality_category'] = 'HD (720p)'
                    elif width >= 854:  # 480p
                        metadata['quality_category'] = 'SD (480p)'
                    else:
                        metadata['quality_category'] = 'Low Resolution'
                
                # Add frame rate if available
                if hasattr(info, 'fps'):
                    fps = getattr(info, 'fps', 0)
                    metadata['fps'] = fps
                    metadata['frame_rate'] = f'{fps} fps' if fps > 0 else 'Unknown'
                
                # Calculate estimated data rate
                if duration > 0 and bitrate > 0:
                    metadata['data_rate_mbps'] = round(bitrate / 1000, 2)
            
            # Reduced tag extraction for performance
            if video_file.tags:
                tags = {}
                # Reduced tag mappings for better performance
                essential_tag_mappings = {
                    'TIT2': 'title', 'TITLE': 'title', '\xa9nam': 'title',
                    'TPE1': 'artist', 'ARTIST': 'artist', '\xa9ART': 'artist',
                    'TDRC': 'year', 'DATE': 'year', '\xa9day': 'year'
                }
                
                tag_count = 0
                max_tags = 10  # Limit number of tags processed
                
                for key, value in video_file.tags.items():
                    if tag_count >= max_tags:
                        break
                    try:
                        mapped_key = essential_tag_mappings.get(key, key.lower())
                        if isinstance(value, list) and value:
                            clean_value = str(value[0]).strip()[:50]  # Reduced length
                        else:
                            clean_value = str(value).strip()[:50]  # Reduced length
                        
                        if clean_value:
                            tags[mapped_key] = clean_value
                            tag_count += 1
                    except (TypeError, AttributeError):
                        continue
                
                if tags:
                    metadata['tags'] = tags
        
        except Exception as e:
            metadata['error'] = f'Failed to extract video metadata: {str(e)}'
        
        return metadata
    
    def extract_document_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract document-specific metadata with enhanced information."""
        metadata = {}
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                metadata.update(self._extract_pdf_metadata(file_path))
                # Add enhanced PDF analysis
                if 'pages' in metadata:
                    page_count = metadata['pages']
                    estimated_words = page_count * 250
                    reading_time_minutes = round(estimated_words / 200, 1)
                    metadata['estimated_reading_time'] = f'{reading_time_minutes} minutes'
                    metadata['estimated_words'] = estimated_words
                    metadata['document_type'] = 'PDF Document'
            elif extension == '.docx':
                metadata.update(self._extract_docx_metadata(file_path))
                # Add enhanced DOCX analysis
                if 'paragraphs' in metadata:
                    metadata['document_type'] = 'Word Document'
            elif extension == '.txt':
                metadata.update(self._extract_text_metadata(file_path))
                # Add enhanced text analysis
                if 'words' in metadata:
                    word_count = metadata['words']
                    metadata['estimated_reading_time'] = f'{round(word_count / 200, 1)} minutes'
                    metadata['document_type'] = 'Text Document'
            else:
                metadata['document_type'] = 'Unknown Document Type'
                metadata['note'] = f'Limited metadata extraction for {extension} files'
        
        except Exception as e:
            metadata['error'] = f'Failed to extract document metadata: {str(e)}'
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from PDF files."""
        metadata = {}
        
        if not PYPDF2_AVAILABLE:
            metadata['error'] = 'PyPDF2 not available for PDF metadata extraction'
            return metadata
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata.update({
                    'pages': len(pdf_reader.pages),
                    'encrypted': pdf_reader.is_encrypted
                })
                
                if pdf_reader.metadata:
                    pdf_metadata = {}
                    for key, value in pdf_reader.metadata.items():
                        clean_key = key.replace('/', '').lower()
                        pdf_metadata[clean_key] = str(value)
                    metadata['pdf_metadata'] = pdf_metadata
        
        except Exception as e:
            metadata['error'] = f'Failed to extract PDF metadata: {str(e)}'
        
        return metadata
    
    def _extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from DOCX files."""
        metadata = {}
        
        if not PYTHON_DOCX_AVAILABLE:
            metadata['error'] = 'python-docx not available for DOCX metadata extraction'
            return metadata
        
        try:
            doc = DocxDocument(file_path)
            
            # Count elements
            paragraphs = len(doc.paragraphs)
            tables = len(doc.tables)
            
            metadata.update({
                'paragraphs': paragraphs,
                'tables': tables
            })
            
            # Core properties
            if doc.core_properties:
                core_props = {}
                props = doc.core_properties
                
                for attr in ['author', 'category', 'comments', 'content_status',
                           'created', 'identifier', 'keywords', 'language',
                           'last_modified_by', 'last_printed', 'modified',
                           'revision', 'subject', 'title', 'version']:
                    value = getattr(props, attr, None)
                    if value is not None:
                        if isinstance(value, datetime):
                            core_props[attr] = value.isoformat()
                        else:
                            core_props[attr] = str(value)
                
                metadata['core_properties'] = core_props
        
        except Exception as e:
            metadata['error'] = f'Failed to extract DOCX metadata: {str(e)}'
        
        return metadata
    
    def _extract_text_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract metadata from text files."""
        metadata = {}
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                
            metadata.update({
                'lines': len(content.splitlines()),
                'words': len(content.split()),
                'characters': len(content),
                'characters_no_spaces': len(content.replace(' ', ''))
            })
        
        except Exception as e:
            metadata['error'] = f'Failed to extract text metadata: {str(e)}'
        
        return metadata
    
    def _categorize_error(self, error: Exception, file_path: Path) -> Dict[str, Any]:
        """Categorize errors and determine severity and recovery options."""
        error_info = {
            'error_type': type(error).__name__,
            'error_message': str(error),
            'severity': ErrorSeverity.MEDIUM.value,
            'recoverable': False,
            'retry_recommended': False
        }
        
        # Categorize based on error type
        if isinstance(error, (FileNotFoundError, PermissionError)):
            error_info.update({
                'severity': ErrorSeverity.HIGH.value,
                'category': 'file_access',
                'recoverable': False
            })
        elif isinstance(error, (OSError, IOError)):
            error_info.update({
                'severity': ErrorSeverity.MEDIUM.value,
                'category': 'io_error',
                'recoverable': True,
                'retry_recommended': True
            })
        elif isinstance(error, MemoryError):
            error_info.update({
                'severity': ErrorSeverity.CRITICAL.value,
                'category': 'memory_error',
                'recoverable': False
            })
        elif isinstance(error, (ImportError, ModuleNotFoundError)):
            error_info.update({
                'severity': ErrorSeverity.HIGH.value,
                'category': 'dependency_missing',
                'recoverable': False
            })
        elif 'corrupt' in str(error).lower() or 'invalid' in str(error).lower():
            error_info.update({
                'severity': ErrorSeverity.HIGH.value,
                'category': 'corrupted_file',
                'recoverable': False
            })
        else:
            error_info.update({
                'severity': ErrorSeverity.MEDIUM.value,
                'category': 'unknown',
                'recoverable': True,
                'retry_recommended': False
            })
        
        return error_info
    
    def _attempt_recovery(self, file_path: Path, error_info: Dict[str, Any]) -> Optional[Dict[str, Any]]:
        """Attempt to recover from extraction errors with fallback methods."""
        if not error_info.get('recoverable', False):
            return None
        
        try:
            # For IO errors, try basic metadata extraction only
            if error_info['category'] == 'io_error':
                logger.info(f"Attempting basic metadata recovery for {file_path}")
                basic_meta = self.extract_basic_metadata(file_path)
                if 'error' not in basic_meta:
                    basic_meta['recovery_method'] = 'basic_only'
                    basic_meta['original_error'] = error_info['error_message']
                    return {'basic': basic_meta}
            
            # For corrupted files, try to extract what we can
            elif error_info['category'] == 'corrupted_file':
                logger.info(f"Attempting partial extraction for potentially corrupted file {file_path}")
                try:
                    stat = file_path.stat()
                    return {
                        'basic': {
                            'filename': file_path.name,
                            'filepath': str(file_path.absolute()),
                            'size': stat.st_size,
                            'size_human': self._format_size(stat.st_size),
                            'extension': file_path.suffix.lower(),
                            'file_type': self.get_file_type(file_path),
                            'recovery_method': 'partial_corrupted',
                            'warning': 'File may be corrupted - limited metadata available'
                        }
                    }
                except Exception:
                    return None
        
        except Exception as recovery_error:
            logger.warning(f"Recovery attempt failed for {file_path}: {recovery_error}")
            return None
        
        return None
    
    def extract_metadata(self, file_path: Union[str, Path], max_retries: int = 2) -> Dict[str, Any]:
        """Extract comprehensive metadata from a file with enhanced error handling and recovery."""
        # Convert string path to Path object if needed
        if isinstance(file_path, str):
            file_path = Path(file_path)
        
        # Basic validation with enhanced error reporting
        try:
            if not file_path.exists():
                raise FileAccessError(f'File not found: {file_path}')
            
            if not file_path.is_file():
                raise FileAccessError(f'Path is not a file: {file_path}')
        
        except FileAccessError as e:
            error_info = self._categorize_error(e, file_path)
            return {
                'error': str(e),
                'filepath': str(file_path),
                'status': 'file_access_error',
                'error_details': error_info,
                'extraction_timestamp': datetime.now().isoformat(),
                'extractor_version': '3.1.0'
            }
        
        # Check cache first
        cache_key = self._get_cache_key(file_path)
        try:
            cached_result = self._get_cached_metadata(cache_key)
            if cached_result is not None:
                cached_result['from_cache'] = True
                return cached_result
        except Exception as cache_error:
            logger.warning(f"Cache access failed for {file_path}: {cache_error}")
        
        # Main extraction with retry logic
        last_error = None
        for attempt in range(max_retries + 1):
            try:
                # Start with basic metadata
                metadata = {'basic': self.extract_basic_metadata(file_path)}
                
                # If basic metadata extraction failed, try recovery
                if 'error' in metadata['basic']:
                    error_info = self._categorize_error(Exception(metadata['basic']['error']), file_path)
                    recovery_result = self._attempt_recovery(file_path, error_info)
                    if recovery_result:
                        return recovery_result
                    return metadata['basic']
                
                file_type = self.get_file_type(file_path)
                metadata['basic']['file_type'] = file_type
                
                # Add extraction metadata
                metadata['basic']['extraction_timestamp'] = datetime.now().isoformat()
                metadata['basic']['extractor_version'] = '3.1.0'
                metadata['basic']['attempt_number'] = attempt + 1
                
                # Extract type-specific metadata with enhanced error handling
                extraction_errors = []
                
                try:
                    if file_type == 'image':
                        image_meta = self.extract_image_metadata(file_path)
                        if image_meta and 'error' not in image_meta:
                            metadata['image'] = image_meta
                        elif 'error' in image_meta:
                            extraction_errors.append(('image', image_meta['error']))
                            
                    elif file_type == 'audio':
                        audio_meta = self.extract_audio_metadata(file_path)
                        if audio_meta and 'error' not in audio_meta:
                            metadata['audio'] = audio_meta
                        elif 'error' in audio_meta:
                            extraction_errors.append(('audio', audio_meta['error']))
                            
                    elif file_type == 'video':
                        video_meta = self.extract_video_metadata(file_path)
                        if video_meta and 'error' not in video_meta:
                            metadata['video'] = video_meta
                        elif 'error' in video_meta:
                            extraction_errors.append(('video', video_meta['error']))
                            
                    elif file_type == 'document':
                        doc_meta = self.extract_document_metadata(file_path)
                        if doc_meta and 'error' not in doc_meta:
                            metadata['document'] = doc_meta
                        elif 'error' in doc_meta:
                            extraction_errors.append(('document', doc_meta['error']))
                            
                except Exception as type_error:
                    extraction_errors.append((file_type, str(type_error)))
                
                # Add extraction errors if any
                if extraction_errors:
                    metadata['basic']['extraction_warnings'] = {
                        f'{error_type}_error': error_msg for error_type, error_msg in extraction_errors
                    }
                
                # Add success status
                metadata['basic']['status'] = 'success' if not extraction_errors else 'partial_success'
                metadata['basic']['from_cache'] = False
                
                # Cache the result (with error handling)
                try:
                    self._cache_metadata(cache_key, metadata)
                except Exception as cache_error:
                    logger.warning(f"Failed to cache result for {file_path}: {cache_error}")
                    metadata['basic']['cache_warning'] = f'Failed to cache result: {str(cache_error)}'
                
                return metadata
                
            except Exception as e:
                last_error = e
                error_info = self._categorize_error(e, file_path)
                
                # Log the error
                logger.warning(f"Extraction attempt {attempt + 1} failed for {file_path}: {e}")
                
                # Try recovery if this is not the last attempt
                if attempt < max_retries and error_info.get('retry_recommended', False):
                    time.sleep(0.1 * (attempt + 1))  # Progressive backoff
                    continue
                
                # Attempt recovery on final failure
                recovery_result = self._attempt_recovery(file_path, error_info)
                if recovery_result:
                    recovery_result['basic']['recovery_used'] = True
                    return recovery_result
                
                # Return detailed error information
                return {
                    'error': f'Failed to extract metadata after {max_retries + 1} attempts: {str(e)}',
                    'filepath': str(file_path),
                    'status': 'extraction_failed',
                    'error_details': error_info,
                    'attempts_made': attempt + 1,
                    'extraction_timestamp': datetime.now().isoformat(),
                    'extractor_version': '3.1.0'
                }
        
        # This should never be reached, but just in case
        error_info = self._categorize_error(last_error or Exception("Unknown error"), file_path)
        return {
            'error': f'Extraction failed after all retry attempts: {str(last_error)}',
            'filepath': str(file_path),
            'status': 'extraction_failed',
            'error_details': error_info,
            'extraction_timestamp': datetime.now().isoformat(),
            'extractor_version': '3.1.0'
        }
    
    def extract_metadata_batch(self, file_paths: List[Union[str, Path]], 
                              progress_callback: Optional[callable] = None,
                              stop_event: Optional[threading.Event] = None) -> Dict[str, Dict[str, Any]]:
        """Extract metadata from multiple files in parallel with improved memory management.
        
        Args:
            file_paths: List of file paths to process
            progress_callback: Optional callback function for progress updates
            stop_event: Optional event to signal early termination
            
        Returns:
            Dictionary mapping file paths to their metadata
        """
        results = {}
        
        def process_file(file_path):
            # Check for stop signal
            if stop_event and stop_event.is_set():
                return None, None
            
            try:
                path_obj = Path(file_path) if isinstance(file_path, str) else file_path
                
                # Use enhanced extraction with retry logic
                metadata = self.extract_metadata(path_obj, max_retries=1)  # Reduced retries for batch
                
                # Add batch processing metadata
                if isinstance(metadata, dict) and 'basic' in metadata:
                    metadata['basic']['batch_processed'] = True
                elif isinstance(metadata, dict):
                    metadata['batch_processed'] = True
                
                return str(path_obj), metadata
                
            except Exception as e:
                # Enhanced error reporting for batch processing
                error_info = {
                    'error': f'Batch processing failed: {str(e)}',
                    'error_type': type(e).__name__,
                    'batch_processed': True,
                    'status': 'batch_processing_failed',
                    'extraction_timestamp': datetime.now().isoformat()
                }
                return str(file_path), error_info
            finally:
                # Force garbage collection for large files
                import gc
                gc.collect()
        
        # Use reusable thread pool for better resource management
        executor = self._get_thread_pool()
        try:
            # Submit all tasks
            future_to_path = {executor.submit(process_file, fp): fp for fp in file_paths}
            
            # Collect results as they complete
            completed = 0
            memory_check_interval = max(1, len(file_paths) // 20)  # Check memory every 5% of files
            
            for future in as_completed(future_to_path):
                if stop_event and stop_event.is_set():
                    # Cancel remaining futures
                    for f in future_to_path:
                        f.cancel()
                    break
                
                try:
                    file_path, metadata = future.result(timeout=30)  # Add timeout
                    if file_path and metadata:
                        results[file_path] = metadata
                except Exception as e:
                    file_path = str(future_to_path[future])
                    results[file_path] = {'error': f'Processing failed: {str(e)}'}
                
                completed += 1
                
                # Periodic memory monitoring and cleanup
                if completed % memory_check_interval == 0:
                    cleanup_result = self.check_and_cleanup_memory()
                    if cleanup_result:
                        logger.info(f"Memory cleanup triggered during batch processing: {cleanup_result['memory_freed_mb']:.1f}MB freed")
                
                if progress_callback:
                    try:
                        progress_callback(completed, len(file_paths))
                    except Exception:
                        pass  # Don't fail if callback fails
                        
        except Exception as e:
            # Handle executor-level errors
            for fp in file_paths[len(results):]:
                results[str(fp)] = {'error': f'Batch processing failed: {str(e)}'}
        
        return results
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of all supported file extensions."""
        extensions = []
        for ext_list in self.supported_types.values():
            extensions.extend(ext_list)
        return sorted(extensions)
    
    def is_supported(self, file_path: Path) -> bool:
        """Check if file type is supported with O(1) lookup."""
        return file_path.suffix.lower() in self._ALL_EXTENSIONS
    
    def _ensure_thread_pool(self) -> None:
        """Ensure thread pool is initialized."""
        with self._thread_pool_lock:
            if self._thread_pool is None:
                self._thread_pool = ThreadPoolExecutor(max_workers=self.max_workers)
    
    def _get_thread_pool(self) -> ThreadPoolExecutor:
        """Get the reusable thread pool."""
        self._ensure_thread_pool()
        return self._thread_pool
    
    @classmethod
    def shutdown_thread_pool(cls) -> None:
        """Shutdown the class-level thread pool."""
        with cls._thread_pool_lock:
            if cls._thread_pool is not None:
                cls._thread_pool.shutdown(wait=True)
                cls._thread_pool = None
    
    def clear_cache(self) -> None:
        """Clear the metadata cache."""
        self._cache.clear()
    
    def get_cache_stats(self) -> Dict[str, Any]:
        """Get cache statistics."""
        return {
            'cache_size': self._cache.size(),
            'cache_enabled': self.enable_cache,
            'cache_max_size': self._cache.max_size,
            'cache_ttl': self._cache.ttl
        }
    
    @staticmethod
    def _format_size(size_bytes: int) -> str:
        """Format file size in human readable format."""
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} PB"
    
    @staticmethod
    def _format_duration(seconds: float) -> str:
        """Format duration in human readable format."""
        if seconds < 60:
            return f"{seconds:.1f}s"
        elif seconds < 3600:
            minutes = int(seconds // 60)
            secs = int(seconds % 60)
            return f"{minutes}:{secs:02d}"
        else:
            hours = int(seconds // 3600)
            minutes = int((seconds % 3600) // 60)
            secs = int(seconds % 60)
            return f"{hours}:{minutes:02d}:{secs:02d}"